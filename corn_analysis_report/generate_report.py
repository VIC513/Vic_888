from pathlib import Path
import numpy as np
import pandas as pd
from docx import Document


def compute_metrics(pred_csv_path: Path, targets: list[str]) -> dict[str, dict[str, float]]:
    df = pd.read_csv(pred_csv_path)
    metrics: dict[str, dict[str, float]] = {}
    for t in targets:
        y = df[f"{t}_True"].to_numpy(dtype=float)
        p = df[f"{t}_Pred"].to_numpy(dtype=float)
        mse = float(np.mean((y - p) ** 2))
        rmse = float(np.sqrt(mse))
        mae = float(np.mean(np.abs(y - p)))
        y_mean = float(np.mean(y))
        sst = float(np.sum((y - y_mean) ** 2))
        sse = float(np.sum((y - p) ** 2))
        r2 = float(1.0 - sse / sst) if sst != 0 else float("nan")
        metrics[t] = {"MSE": mse, "RMSE": rmse, "MAE": mae, "R2": r2}
    return metrics


def add_metrics_table(doc: Document, metrics: dict[str, dict[str, float]], targets: list[str]) -> None:
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "指标"
    hdr[1].text = "Moisture"
    hdr[2].text = "Oil"
    hdr[3].text = "Protein"
    hdr[4].text = "Starch"

    for m in ["MSE", "RMSE", "MAE", "R2"]:
        row = table.add_row().cells
        row[0].text = m
        for i, t in enumerate(targets, start=1):
            v = metrics[t][m]
            row[i].text = "NaN" if np.isnan(v) else f"{v:.4f}"


def main() -> None:
    base_dir = Path(__file__).resolve().parent
    results_dir = base_dir / "Results"
    pred_csv = results_dir / "Test_Predictions.csv"
    code_file = base_dir / "Assignment.py"
    loss_fig = results_dir / "Loss_Curve.png"
    scatter_fig = results_dir / "Prediction_vs_True.png"
    spectra_fig = results_dir / "nir_spectra_samples.png"

    if not pred_csv.exists():
        raise FileNotFoundError(f"未找到预测结果文件：{pred_csv}。请先运行 Assignment.py 生成 Results/Test_Predictions.csv")

    targets = ["Moisture", "Oil", "Protein", "Starch"]
    metrics = compute_metrics(pred_csv, targets)

    doc = Document()
    doc.add_heading("近红外光谱（NIR）玉米成分预测分析报告", level=0)

    doc.add_heading("一、数据预处理与描述", level=1)
    doc.add_heading("1.1 数据集概述", level=2)
    doc.add_paragraph("本研究使用包含 80 个玉米样本的近红外（NIR）光谱数据集进行建模分析。数据由两部分组成：")
    doc.add_paragraph("光谱特征：波长范围 1100–2498 nm，采样间隔 2 nm，共计 700 个通道/特征（对应每个样本的 700 维输入向量）。", style="List Bullet")
    doc.add_paragraph("预测目标（多输出回归）：每个样本对应四个化学成分含量：水分（Moisture）、油脂（Oil）、蛋白质（Protein）和淀粉（Starch）。", style="List Bullet")

    doc.add_heading("1.2 预处理流程", level=2)
    doc.add_paragraph("为提升训练稳定性并减少不同通道尺度差异带来的影响，采用如下预处理：")
    doc.add_paragraph("特征标准化（Z-score）：对 700 维光谱特征使用 StandardScaler 标准化，使每个通道在训练集上满足均值为 0、方差为 1，并将同样的变换应用于测试集。", style="List Bullet")
    doc.add_paragraph("数据集划分：按照 7:3 比例随机划分，训练集 56 个样本，测试集 24 个样本（random_state=42）。", style="List Bullet")
    doc.add_paragraph("补充说明：当前实现仅对输入 X 进行了标准化，未对输出 y（四个成分）做标准化，这会导致不同成分量纲差异对损失函数的权重产生影响。")

    doc.add_heading("二、模型设计", level=1)
    doc.add_paragraph("针对光谱数据“高维（700维）+ 小样本（80条）”特性，设计了基于 PyTorch 的多目标回归深度学习模型（全连接神经网络 FNN）。")
    doc.add_heading("2.1 网络架构", level=2)
    doc.add_paragraph("模型采用“共享特征提取 + 多输出回归”的结构：")
    doc.add_paragraph("输入层：700 维光谱向量。", style="List Bullet")
    doc.add_paragraph("隐藏层：2 层全连接网络，并逐层减半隐藏维度（700→256→128）。", style="List Bullet")
    doc.add_paragraph("输出层：128→4，分别对应 Moisture、Oil、Protein、Starch。", style="List Bullet")
    doc.add_paragraph("隐藏层内包含 Linear、BatchNorm1d、ReLU 与 Dropout，用于增强非线性表达能力并提升训练稳定性。")

    doc.add_heading("2.2 防过拟合策略", level=2)
    doc.add_paragraph("考虑样本量较小，模型引入多种正则化与训练策略：")
    doc.add_paragraph("Dropout：dropout_rate=0.3。", style="List Bullet")
    doc.add_paragraph("L2 正则化：优化器 weight_decay=1e-4。", style="List Bullet")
    doc.add_paragraph("学习率调度：ReduceLROnPlateau 根据评估集损失停滞自动降低学习率。", style="List Bullet")

    doc.add_heading("三、模型训练与实现细节", level=1)
    doc.add_heading("3.1 训练配置", level=2)
    doc.add_paragraph("优化器：Adam（lr=0.005，weight_decay=1e-4）。", style="List Bullet")
    doc.add_paragraph("损失函数：MSELoss。", style="List Bullet")
    doc.add_paragraph("训练轮数：500 epochs；Batch size：16。", style="List Bullet")
    doc.add_paragraph("Checkpoint：每 100 个 epoch 保存一次模型参数（CheckPoints/）。", style="List Bullet")

    doc.add_heading("3.2 训练监控", level=2)
    doc.add_paragraph("记录 Train Loss 与评估集 Loss 的变化趋势，并输出损失曲线图（Results/Loss_Curve.png）。")
    doc.add_paragraph("说明：当前实现中训练过程中使用测试集计算评估损失，这会造成一定程度的信息泄露。更严格的做法应再划分验证集，或使用 K 折交叉验证。")

    doc.add_heading("四、模型评估与结果分析", level=1)
    doc.add_heading("4.1 性能指标", level=2)
    doc.add_paragraph("在测试集（24 条样本）上分别对四个成分计算 MSE、RMSE、MAE 与 R²。指标来自 Results/Test_Predictions.csv 的真实值与预测值。")
    add_metrics_table(doc, metrics, targets)
    doc.add_paragraph("结果解读：当 R² 为负时，说明模型在该测试划分下的泛化效果未达到理想水平，可能不如以训练集均值作为常数预测的简单基线。")

    doc.add_heading("4.2 预测效果可视化", level=2)
    doc.add_paragraph("报告生成了“真实值 vs 预测值”的散点对比图（Results/Prediction_vs_True.png）。理想情况下散点应沿 y=x 附近分布。")

    doc.add_heading("五、总结与改进方案", level=1)
    doc.add_heading("5.1 当前方案的优点", level=2)
    doc.add_paragraph("端到端流程完整：包含数据读取、标准化、训练、保存模型、生成预测结果文件与可视化图表。", style="List Bullet")
    doc.add_paragraph("多输出建模：单模型同时预测四个成分，具备共享特征学习能力。", style="List Bullet")
    doc.add_paragraph("具备基本正则化手段：Dropout、L2 正则、学习率调度等。", style="List Bullet")

    doc.add_heading("5.2 改进建议", level=2)
    doc.add_paragraph("目标值标准化：对 y 也做标准化（或分别 z-score），推理阶段再反标准化，以缓解不同量纲对 MSE 的影响。", style="List Bullet")
    doc.add_paragraph("严格验证策略：从训练集再拆分出验证集（train/val/test），或采用 K 折交叉验证，避免训练过程反复“看见测试集”。", style="List Bullet")
    doc.add_paragraph("损失加权或分目标建模：对不同成分设置损失权重，或分别训练四个单输出模型作基线对比。", style="List Bullet")
    doc.add_paragraph("引入光谱建模基线：PLSR/SVR 作为 baseline，以便更客观地评估深度学习方案的收益。", style="List Bullet")
    doc.add_paragraph("采用更适合光谱的结构与预处理：SNV/MSC、Savitzky–Golay 平滑与导数；以及 1D-CNN/TCN 等利用谱序列局部相关性的模型。", style="List Bullet")

    doc.add_heading("提交附件说明", level=1)
    doc.add_paragraph(f"代码：{code_file.name}", style="List Bullet")
    doc.add_paragraph(f"结果：{pred_csv.name}", style="List Bullet")
    doc.add_paragraph("图表：Results/ 目录下的 Loss 曲线图、散点对比图与光谱示例图。", style="List Bullet")
    doc.add_paragraph(f"Loss 曲线图：{loss_fig.name}", style="List Bullet")
    doc.add_paragraph(f"散点对比图：{scatter_fig.name}", style="List Bullet")
    doc.add_paragraph(f"光谱示例图：{spectra_fig.name}", style="List Bullet")

    out_path = base_dir / "近红外光谱（NIR）玉米成分预测分析报告.docx"
    doc.save(out_path)
    print(str(out_path))


if __name__ == "__main__":
    main()

